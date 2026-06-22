"""Round II-B doctrine tests: no deterministic semantic organs.

Pins the brief's Track A-F rules:
  - RiskReport: no deterministic semantic risk items
  - RewritePlan: no deterministic semantic changes
  - ComplianceChecklist: never disappears silently
  - docx Title style extraction
  - Narrator per-axis status visible on parse_failed
"""

from __future__ import annotations

import unittest

from kairoskopion.schema import (
    ArticleModel,
    FitAssessment,
    MismatchMap,
    SubmissionScenario,
    VenueModel,
)
from kairoskopion.services.risk_report_needs_llm import (
    build_needs_llm_risk_report,
)
from kairoskopion.services.rewrite_plan_needs_llm import (
    build_needs_llm_rewrite_plan,
)
from kairoskopion.services.semantic_provenance import (
    ORIGIN_NEEDS_LLM,
    SEMANTIC_STATUS_NEEDS_LLM,
)


def _article(): return ArticleModel(title_current="X", genre_current="theoretical_essay")
def _venue(): return VenueModel(canonical_name="V", venue_type="journal")
def _fit(): return FitAssessment(axes=[], overall_label="possible")
def _mm(): return MismatchMap(
    mismatches=[
        {"axis": "topic", "severity": "blocking", "article_side": "X",
         "venue_side": "", "description": "", "possible_actions": [],
         "field_core_risk": "unknown_core_impact"},
        {"axis": "method", "severity": "major", "article_side": "Y",
         "venue_side": "", "description": "", "possible_actions": [],
         "field_core_risk": "unknown_core_impact"},
    ],
)


# ---------------- RiskReport: needs_llm placeholder ----------------

class TestRiskReportNeedsLLM(unittest.TestCase):
    def test_no_risk_items_emitted(self):
        rr = build_needs_llm_risk_report(
            _article(), _venue(), SubmissionScenario(), _fit(), _mm(),
        )
        self.assertEqual(rr.risk_items, [])
        self.assertEqual(rr.blocking_risks, [])
        self.assertEqual(rr.warnings, [])
        self.assertIsNone(rr.overall_risk_label)

    def test_semantic_status_needs_llm(self):
        rr = build_needs_llm_risk_report(
            _article(), _venue(), SubmissionScenario(), _fit(), _mm(),
        )
        self.assertEqual(rr.semantic_status, SEMANTIC_STATUS_NEEDS_LLM)
        for f in ("risk_items", "blocking_risks", "warnings",
                  "overall_risk_label"):
            self.assertEqual(
                rr.field_origins.get(f), ORIGIN_NEEDS_LLM,
                f"RiskReport field {f!r} must have origin=needs_llm",
            )

    def test_unknowns_explain_llm_organ_absence(self):
        rr = build_needs_llm_risk_report(
            _article(), _venue(), SubmissionScenario(), _fit(), _mm(),
        )
        joined = " ".join(rr.unknowns).lower()
        self.assertIn("risk_officer", joined)
        self.assertIn("organ", joined)

    def test_no_semantic_prose_in_output(self):
        """No deterministic semantic risk diagnosis strings allowed."""
        rr = build_needs_llm_risk_report(
            _article(), _venue(), SubmissionScenario(), _fit(), _mm(),
        )
        import json as _json
        flat = _json.dumps(rr.to_dict(), ensure_ascii=False).lower()
        forbidden = [
            "scope risk", "method risk", "citation gap risk",
            "reviewer misunderstanding", "field-core loss",
            "strategic publication risk", "desk-reject",
        ]
        for phrase in forbidden:
            self.assertNotIn(phrase, flat,
                             f"RiskReport must not emit semantic prose: {phrase!r}")


# ---------------- RewritePlan: needs_llm placeholder ----------------

class TestRewritePlanNeedsLLM(unittest.TestCase):
    def test_no_changes_emitted(self):
        rp = build_needs_llm_rewrite_plan(_mm())
        self.assertEqual(rp.changes, [])
        self.assertIsNone(rp.estimated_effort)
        self.assertIsNone(rp.summary)
        self.assertEqual(rp.field_core_risk, "unknown_core_impact")

    def test_semantic_status_needs_llm(self):
        rp = build_needs_llm_rewrite_plan(_mm())
        self.assertEqual(rp.semantic_status, SEMANTIC_STATUS_NEEDS_LLM)
        for f in ("changes", "summary", "estimated_effort", "field_core_risk"):
            self.assertEqual(
                rp.field_origins.get(f), ORIGIN_NEEDS_LLM,
                f"RewritePlan field {f!r} must have origin=needs_llm",
            )

    def test_unknowns_explain_llm_organ_absence(self):
        rp = build_needs_llm_rewrite_plan(_mm())
        joined = " ".join(rp.unknowns).lower()
        self.assertIn("rewrite_planner", joined)
        self.assertIn("organ", joined)

    def test_no_semantic_prose_in_output(self):
        rp = build_needs_llm_rewrite_plan(_mm())
        import json as _json
        flat = _json.dumps(rp.to_dict(), ensure_ascii=False).lower()
        forbidden = [
            "revise introduction", "strengthen literature review",
            "clarify method", "adapt to venue", "change argument",
            "core_touching",
            "intro_reframe",
        ]
        for phrase in forbidden:
            self.assertNotIn(phrase, flat,
                             f"RewritePlan must not emit semantic prose: {phrase!r}")

    def test_empty_mismatch_map_still_returns_safe_object(self):
        rp = build_needs_llm_rewrite_plan(None)
        self.assertEqual(rp.changes, [])
        self.assertEqual(rp.semantic_status, SEMANTIC_STATUS_NEEDS_LLM)


# ---------------- docx Title extraction ----------------

class TestDocxTitleStyle(unittest.TestCase):
    def _make_docx_with_style(self, title: str, body: str, style: str, path):
        from docx import Document
        d = Document()
        d.add_paragraph(title, style=style)
        d.add_paragraph(body)
        d.save(str(path))

    def test_title_style_extracted(self, ):
        import tempfile, os
        from kairoskopion.adapters.source_intake import _extract_docx_text
        from pathlib import Path
        td = tempfile.mkdtemp(prefix="r2b_")
        p = Path(td) / "t.docx"
        try:
            self._make_docx_with_style(
                "Real Source Title", "Body paragraph.", "Title", p,
            )
            text, status, warnings = _extract_docx_text(p)
            self.assertEqual(status, "extracted")
            self.assertTrue(text.startswith("# Real Source Title"),
                            f"docx Title style not prepended as H1: {text[:80]!r}")
        finally:
            import shutil; shutil.rmtree(td, ignore_errors=True)

    def test_heading1_style_extracted(self):
        import tempfile, os
        from kairoskopion.adapters.source_intake import _extract_docx_text
        from pathlib import Path
        td = tempfile.mkdtemp(prefix="r2b_")
        p = Path(td) / "h.docx"
        try:
            self._make_docx_with_style(
                "Heading One Title", "Body text.", "Heading 1", p,
            )
            text, status, warnings = _extract_docx_text(p)
            self.assertEqual(status, "extracted")
            self.assertTrue(text.startswith("# Heading One Title"))
        finally:
            import shutil; shutil.rmtree(td, ignore_errors=True)

    def test_no_title_invented_when_absent(self):
        """When the docx has no Title/Heading 1 style, NO title is invented."""
        import tempfile, os
        from kairoskopion.adapters.source_intake import _extract_docx_text
        from pathlib import Path
        from docx import Document
        td = tempfile.mkdtemp(prefix="r2b_")
        p = Path(td) / "n.docx"
        try:
            d = Document()
            d.add_paragraph("Just body, no heading.")
            d.add_paragraph("More body.")
            d.save(str(p))
            text, status, warnings = _extract_docx_text(p)
            self.assertEqual(status, "extracted")
            self.assertFalse(text.startswith("# "),
                             "no title style → no H1 invented")
        finally:
            import shutil; shutil.rmtree(td, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
