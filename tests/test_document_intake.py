"""Tests for Sprint 1: Real Document Intake (PDF, DOCX, extraction statuses)."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from kairoskopion.adapters.source_intake import (
    EXTRACTION_STATUSES,
    SourceRole,
    register_local_source,
    register_text_input,
    _content_hash,
    _extract_html_text,
)


# ---------------------------------------------------------------------------
# Text-based intake (existing behaviour preserved)
# ---------------------------------------------------------------------------

class TestTextIntake:
    def test_markdown_extraction(self, tmp_path):
        md = tmp_path / "paper.md"
        md.write_text("# Title\n\nBody text here.", encoding="utf-8")
        snap, text = register_local_source(md, role=SourceRole.ARTICLE_INPUT)
        assert snap.extraction_status == "extracted"
        assert snap.parser_used == "local_file_read"
        assert "Title" in text
        assert snap.content_hash is not None

    def test_txt_extraction(self, tmp_path):
        txt = tmp_path / "notes.txt"
        txt.write_text("Plain text notes.", encoding="utf-8")
        snap, text = register_local_source(txt)
        assert snap.extraction_status == "extracted"
        assert text == "Plain text notes."

    def test_json_extraction(self, tmp_path):
        j = tmp_path / "scenario.json"
        j.write_text(json.dumps({"goal": "test"}), encoding="utf-8")
        snap, text = register_local_source(j, role=SourceRole.SUBMISSION_INFO)
        assert snap.extraction_status == "extracted"
        assert "goal" in text

    def test_html_extraction(self, tmp_path):
        h = tmp_path / "page.html"
        h.write_text("<html><body><p>Hello</p><script>var x=1;</script></body></html>",
                      encoding="utf-8")
        snap, text = register_local_source(h)
        assert snap.extraction_status == "extracted"
        assert "Hello" in text
        assert "var x" not in text  # script stripped
        assert snap.parser_used == "html_tag_strip"


# ---------------------------------------------------------------------------
# PDF intake
# ---------------------------------------------------------------------------

class TestPDFIntake:
    def _make_simple_pdf(self, path: Path) -> None:
        """Create a minimal valid PDF with text using pypdf."""
        try:
            from pypdf import PdfWriter
            from pypdf.generic import NameObject, TextStringObject
        except ImportError:
            pytest.skip("pypdf not installed")

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        # Add a simple page with metadata (text extraction from blank page is empty,
        # so we test the extraction flow, not rich text)
        writer.add_metadata({"/Title": "Test PDF"})
        with open(path, "wb") as f:
            writer.write(f)

    def test_pdf_extraction_runs(self, tmp_path):
        """PDF extraction runs without error (may return empty if blank page)."""
        pdf = tmp_path / "paper.pdf"
        self._make_simple_pdf(pdf)
        snap, text = register_local_source(pdf, role=SourceRole.ARTICLE_INPUT)
        assert snap.extraction_status in ("extracted", "needs_ocr", "partially_extracted")
        assert snap.parser_used == "pypdf"
        assert snap.content_type == "application/pdf"

    def test_missing_pdf(self, tmp_path):
        snap, text = register_local_source(tmp_path / "missing.pdf")
        assert snap.extraction_status == "file_not_found"
        assert text == ""

    def test_corrupted_pdf(self, tmp_path):
        pdf = tmp_path / "bad.pdf"
        pdf.write_bytes(b"not a pdf")
        snap, text = register_local_source(pdf)
        assert snap.extraction_status in ("encrypted_or_unreadable", "failed")
        assert text == ""


# ---------------------------------------------------------------------------
# DOCX intake
# ---------------------------------------------------------------------------

class TestDOCXIntake:
    def _make_simple_docx(self, path: Path) -> None:
        """Create a minimal valid DOCX with text."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is the body text of the test document.")
        doc.add_paragraph("Second paragraph with more content.")
        doc.save(str(path))

    def test_docx_extraction(self, tmp_path):
        docx = tmp_path / "paper.docx"
        self._make_simple_docx(docx)
        snap, text = register_local_source(docx, role=SourceRole.ARTICLE_INPUT)
        assert snap.extraction_status == "extracted"
        assert snap.parser_used == "python-docx"
        assert "body text" in text
        assert snap.content_hash is not None

    def test_corrupted_docx(self, tmp_path):
        docx = tmp_path / "bad.docx"
        docx.write_bytes(b"not a docx")
        snap, text = register_local_source(docx)
        assert snap.extraction_status == "failed"
        assert text == ""

    def test_empty_docx(self, tmp_path):
        """DOCX with no paragraphs."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx))
        snap, text = register_local_source(docx)
        assert snap.extraction_status == "partially_extracted"


# ---------------------------------------------------------------------------
# Unsupported / binary
# ---------------------------------------------------------------------------

class TestUnsupportedFormats:
    def test_unsupported_extension(self, tmp_path):
        f = tmp_path / "data.xlsx"
        f.write_bytes(b"PK binary")
        snap, text = register_local_source(f)
        assert snap.extraction_status == "unsupported"
        assert text == ""
        assert snap.extraction_errors

    def test_binary_file(self, tmp_path):
        f = tmp_path / "image.png"
        f.write_bytes(b"\x89PNG\r\n\x1a\n")
        snap, text = register_local_source(f)
        assert snap.extraction_status == "unsupported"
        assert text == ""


# ---------------------------------------------------------------------------
# Content hash stability
# ---------------------------------------------------------------------------

class TestContentHash:
    def test_hash_stable(self, tmp_path):
        f = tmp_path / "doc.md"
        content = "Stable content for hashing"
        f.write_text(content, encoding="utf-8")
        s1, _ = register_local_source(f)
        s2, _ = register_local_source(f)
        assert s1.content_hash == s2.content_hash

    def test_hash_changes_with_content(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("version 1", encoding="utf-8")
        s1, _ = register_local_source(f)
        f.write_text("version 2", encoding="utf-8")
        s2, _ = register_local_source(f)
        assert s1.content_hash != s2.content_hash


# ---------------------------------------------------------------------------
# Metadata and source snapshot fields
# ---------------------------------------------------------------------------

class TestMetadata:
    def test_snapshot_has_required_fields(self, tmp_path):
        f = tmp_path / "test.md"
        f.write_text("content", encoding="utf-8")
        snap, text = register_local_source(f, role=SourceRole.ARTICLE_INPUT)
        assert snap.snapshot_id
        assert snap.source_id
        assert snap.url  # file path
        assert snap.retrieved_at
        assert snap.content_type == "text/markdown"
        assert snap.extraction_status in EXTRACTION_STATUSES

    def test_source_id_default(self, tmp_path):
        f = tmp_path / "myfile.md"
        f.write_text("x", encoding="utf-8")
        snap, _ = register_local_source(f)
        assert snap.source_id == "local:myfile.md"

    def test_source_id_override(self, tmp_path):
        f = tmp_path / "myfile.md"
        f.write_text("x", encoding="utf-8")
        snap, _ = register_local_source(f, source_id="custom:123")
        assert snap.source_id == "custom:123"


# ---------------------------------------------------------------------------
# HTML tag stripping
# ---------------------------------------------------------------------------

class TestHTMLStripping:
    def test_strips_tags(self):
        assert "Hello" in _extract_html_text("<p>Hello</p>")

    def test_strips_scripts(self):
        html = "<p>Text</p><script>alert(1);</script><p>More</p>"
        result = _extract_html_text(html)
        assert "alert" not in result
        assert "Text" in result
        assert "More" in result

    def test_strips_style(self):
        html = "<style>.x{color:red}</style><p>Content</p>"
        result = _extract_html_text(html)
        assert "color" not in result
        assert "Content" in result


# ---------------------------------------------------------------------------
# Extraction status taxonomy
# ---------------------------------------------------------------------------

class TestExtractionStatuses:
    def test_all_statuses_are_strings(self):
        for s in EXTRACTION_STATUSES:
            assert isinstance(s, str)

    def test_required_statuses_present(self):
        required = {"extracted", "partially_extracted", "unsupported", "failed",
                     "binary_not_extracted", "needs_ocr", "encrypted_or_unreadable",
                     "file_not_found", "unknown"}
        assert required.issubset(set(EXTRACTION_STATUSES))
